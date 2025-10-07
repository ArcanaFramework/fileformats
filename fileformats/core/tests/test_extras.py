import platform
import typing as ty
from pathlib import Path

import pytest
from docx import Document

from fileformats.core import FileSet, MockMixin, extra, extra_implementation
from fileformats.testing import Foo
from fileformats.vendor.openxmlformats_officedocument.application import (
    Wordprocessingml_Document,
)


def test_sample():
    test_inst = Foo.sample()
    assert test_inst.fspath.exists()
    assert test_inst.fspath.suffix == ".foo"


def test_mock():
    mock = Foo.mock()
    if platform.system() == "Windows":
        expected = Path(f"{Path().cwd().drive}\\mock\\foo.foo")
    else:
        expected = Path("/mock/foo.foo")
    assert mock.fspath == expected
    assert not mock.fspath.exists()
    assert isinstance(mock, MockMixin)


class Woo(FileSet):
    @extra
    def test_extra(self, a: int, b: float, c: ty.Optional[str] = None) -> float:
        raise NotImplementedError


def test_extra_signature_no_default():
    extra_implementation(Woo.test_extra)

    def woo_test_extra(woo: Woo, a: int, b: float) -> float:
        pass


def test_extra_signature1():

    with pytest.raises(TypeError, match="missing required argument"):

        @extra_implementation(Woo.test_extra)
        def woo_test_extra(woo: Woo, a: int) -> float:
            pass


def test_extra_signature2():

    with pytest.raises(TypeError, match="name of parameter"):

        @extra_implementation(Woo.test_extra)
        def woo_test_extra(woo: Woo, a: int, d: str) -> float:
            pass


def test_extra_signature3():

    with pytest.raises(TypeError, match="found additional argument"):

        @extra_implementation(Woo.test_extra)
        def woo_test_extra(
            woo: Woo, a: int, b: float, c: ty.Optional[str], d: int
        ) -> float:
            pass


def test_extra_signature4():

    with pytest.raises(TypeError, match="return type"):

        @extra_implementation(Woo.test_extra)
        def woo_test_extra(woo: Woo, a: int, b: str) -> int:
            pass


def test_vendor_extra_load(tmp_path: Path):

    fspath = tmp_path / "test.docx"

    # Create a sample document
    data = Document()
    data.add_heading("A document")
    data.add_paragraph("The quick brown fox jumped over the lazy dog.")

    Wordprocessingml_Document.new(fspath, data=data)
    doc = Wordprocessingml_Document(fspath)
    hsh = doc.hash()
    reloaded_data = doc.load()
    assert documents_equal(data, reloaded_data)
    data.add_paragraph("Another paragraph")
    doc.save(data)
    assert doc.hash() != hsh


def documents_equal(doc1: Document, doc2: Document) -> bool:
    """Compare two Document objects by their XML content."""

    # Get the main document part XML
    xml1 = doc1.part.element.xml
    xml2 = doc2.part.element.xml

    return xml1 == xml2
